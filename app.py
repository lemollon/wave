# app.py — Wave (Pulse companion) — Trends → Leads → Outreach → Weekly Report + Pro Lab
# Adds: clean AI suggestions, clickable chips, persistent lists, global Industry banner, Clear-all context.

import os
import io
import re
import json
import time
import textwrap
import datetime as dt
from typing import Dict, List, Optional, Tuple, Any

import streamlit as st
import pandas as pd
import requests

# ---------- Optional libs with graceful fallbacks ----------
DOCX_OK = True
try:
    from docx import Document
except Exception:
    DOCX_OK = False

MAPS_OK = True
try:
    import folium
    from streamlit_folium import st_folium
except Exception:
    MAPS_OK = False

# OpenAI v1 SDK import
try:
    from openai import OpenAI
except Exception:
    OpenAI = None

# Pro libs (optional)
LC_OK = LG_OK = LCO_OK = False
try:
    from langchain_core.prompts import ChatPromptTemplate
    LC_OK = True
except Exception:
    LC_OK = False

try:
    from langgraph.graph import StateGraph, END
    LG_OK = True
except Exception:
    LG_OK = False

try:
    from langchain_openai import ChatOpenAI
    LCO_OK = True
except Exception:
    LCO_OK = False

CREW_OK = False
try:
    from crewai import Agent, Task, Crew
    CREW_OK = True
except Exception:
    CREW_OK = False


# -------------------- ENV / helpers --------------------
def _env(k: str, d: str = "") -> str:
    """
    Look up a value from OS env; if absent, fall back to Streamlit secrets.
    Works on Render (env vars) and Streamlit Cloud (secrets).
    """
    v = os.getenv(k)
    if not v:
        try:
            v = st.secrets.get(k, d)  # safe even if secrets isn't configured
        except Exception:
            v = d
    return v


OPENAI_API_KEY = _env("OPENAI_API_KEY", "")

# Safe client init for Streamlit + Render
if OpenAI and OPENAI_API_KEY:
    try:
        client = OpenAI(api_key=OPENAI_API_KEY)
    except Exception as _e:
        client = None
else:
    client = None


def llm_ok() -> bool:
    return client is not None


def llm(prompt: str, system: str = "You are a helpful marketer.", temp: float = 0.4) -> str:
    """Small wrapper around OpenAI chat; returns '' if not configured."""
    if not llm_ok():
        return ""
    try:
        r = client.chat.completions.create(
            model="gpt-4o-mini",
            temperature=temp,
            messages=[
                {"role": "system", "content": system},
                {"role": "user", "content": prompt},
            ],
        )
        return (r.choices[0].message.content or "").strip()
    except Exception as e:
        return f"(AI unavailable: {e})"


def build_docx_bytes(title: str, body_md: str) -> bytes:
    if not DOCX_OK:
        return b""
    doc = Document()
    doc.add_heading(title, level=1)
    for para in body_md.split("\n\n"):
        if para.strip():
            doc.add_paragraph(para.strip())
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# -------------------- UI helpers (NEW) --------------------
def inject_css():
    st.markdown(
        """
        <style>
          :root { --card-bg:#0e1117; --card-border:#2b2f36; }
          .kpi-card {border:1px solid var(--card-border); border-radius:14px; padding:16px 18px; background:var(--card-bg);}
          .kpi-label {font-size:.80rem; opacity:.85; letter-spacing:.2px}
          .kpi-value {font-weight:800; font-size:1.25rem; margin-top:4px}
          .stDataFrame { border:1px solid var(--card-border); border-radius:12px; }
          .stAlert { border-radius:12px; }
          .stButton>button { border-radius:10px; padding:.5rem .9rem; font-weight:600 }
          .chip { display:inline-block; padding:.35rem .6rem; border:1px solid #2b2f36; border-radius:999px; margin:.25rem .35rem .25rem 0; background:#121621; font-size:.9rem; cursor:pointer;}
          .chip:hover { background:#1a1f27; }
          .helper-banner {
            border:1px dashed #2b2f36; border-radius:10px; padding:.6rem .8rem; font-size:.9rem; margin: .5rem 0 1rem 0;
            background:rgba(255,255,255,0.02);
          }
        </style>
        """,
        unsafe_allow_html=True,
    )


def kpi(label: str, value: str):
    st.markdown(
        f"""
        <div class="kpi-card">
          <div class="kpi-label">{label}</div>
          <div class="kpi-value">{value}</div>
        </div>
        """,
        unsafe_allow_html=True,
    )


def helper_banner():
    st.markdown(
        """
        <div class="helper-banner">
          <b>Tip:</b> Start by choosing your <b>Industry</b> on <i>Trend Rider</i>. That selection powers Lead Finder, Outreach, and Weekly Report.
          You can combine industries (comma-separated). Want a clean slate? Use <b>Sidebar → Clear all context</b>.
        </div>
        """,
        unsafe_allow_html=True,
    )


# ---------- Suggestion sanitation (NEW) ----------
_JSON_FENCE_RE = re.compile(r"^```+|```+$", flags=re.MULTILINE)
_QUOTES_RE = re.compile(r"(^[\"'`\\s]+|[\"'`\\s]+$)")

def extract_strings(raw: Any, max_items: int = 24) -> List[str]:
    """
    Accepts LLM output that might be:
     - a proper JSON array string
     - a dict with a key (rare)
     - plain text with code fences, 'json' labels, brackets, quotes
    Returns a clean list of strings, trimmed & de-duped.
    """
    out: List[str] = []

    if raw is None:
        return out

    if isinstance(raw, (list, tuple)):
        candidates = list(raw)
    else:
        txt = str(raw).strip()

        # remove code fences
        txt = _JSON_FENCE_RE.sub("", txt).strip()

        # if it looks like "json\n[...]" or startswith "json", drop leading label
        if txt.lower().startswith("json"):
            # remove a leading 'json' word and possible colon
            txt = re.sub(r"^json\\s*[:\\-]*", "", txt, flags=re.IGNORECASE).strip()

        # try JSON parse first
        try:
            data = json.loads(txt)
            if isinstance(data, list):
                candidates = data
            elif isinstance(data, dict):
                # if dict: collect string values
                candidates = []
                for v in data.values():
                    if isinstance(v, list):
                        candidates.extend(v)
                    else:
                        candidates.append(v)
            else:
                candidates = [txt]
        except Exception:
            # fallback: strip brackets and split by comma/newline
            txt2 = txt.strip()
            # remove outer brackets if present
            if txt2.startswith("[") and txt2.endswith("]"):
                txt2 = txt2[1:-1]
            # split on commas & newlines
            parts = re.split(r"[\\n,]", txt2)
            candidates = [p for p in parts]

    # clean each token
    for c in candidates:
        s = str(c)
        s = _QUOTES_RE.sub("", s.strip())
        s = s.strip("[]")
        s = s.strip()
        if s:
            out.append(s)

    # de-dupe, preserve order
    seen = set()
    clean = []
    for x in out:
        if x not in seen:
            seen.add(x)
            clean.append(x)

    return clean[:max_items]


def render_chips_and_capture(label: str, suggestions: List[str], key_prefix: str,
                             state_list_name: str, text_input_key: str):
    """
    Renders clickable chips for suggestions. When clicked, item is added to
    a session_state list and sync'd into a comma-separated text input.
    """
    if not suggestions:
        return

    st.markdown(f"**{label}**")
    cols = st.columns(min(4, max(1, len(suggestions)//6 + 1)))
    # display in columns to wrap nicer
    i = 0
    for s in suggestions:
        col = cols[i % len(cols)]
        with col:
            if st.button(s, key=f"{key_prefix}_chip_{i}", use_container_width=False, help="Click to add"):
                lst: List[str] = st.session_state.get(state_list_name, [])
                if s not in lst:
                    lst.append(s)
                    st.session_state[state_list_name] = lst
                # also reflect into text input
                st.session_state[text_input_key] = ", ".join(lst)
        i += 1

    # bulk actions
    b1, b2, _ = st.columns([1,1,6])
    with b1:
        if st.button("Add all", key=f"{key_prefix}_addall"):
            lst = st.session_state.get(state_list_name, [])
            for s in suggestions:
                if s not in lst:
                    lst.append(s)
            st.session_state[state_list_name] = lst
            st.session_state[text_input_key] = ", ".join(lst)
    with b2:
        if st.button("Clear", key=f"{key_prefix}_clear"):
            st.session_state[state_list_name] = []
            st.session_state[text_input_key] = ""


# ------------- Warm-up button (Render/Cloud cold starts) -------------
def warm_up_render(url: str, timeout: int = 10) -> str:
    if not url:
        return "No wake URL set. Add RENDER_WAKE_URL to your environment."
    try:
        r = requests.head(url, timeout=timeout, allow_redirects=True)
        if 200 <= r.status_code < 400:
            return f"Warmed! ({r.status_code})"
        r = requests.get(url, timeout=timeout)
        return f"Warmed! ({r.status_code})" if 200 <= r.status_code < 400 else f"Service responded: {r.status_code}"
    except Exception as e:
        return f"Warm-up failed: {e}"


# --------------------- Trends (inline) --------------------
def google_trends_rising(keywords: List[str], geo="US", timeframe="now 7-d") -> Dict:
    try:
        from pytrends.request import TrendReq
    except Exception:
        return {"source": "google_trends", "error": "pytrends not installed", "rising": [], "iot": {}}
    pytrends = TrendReq(hl="en-US", tz=360)
    rising_all, iot_map = [], {}
    for kw in keywords[:6]:
        try:
            pytrends.build_payload([kw], timeframe=timeframe, geo=geo)
            iot = pytrends.interest_over_time()
            if not iot.empty:
                ser = iot[kw].reset_index().rename(columns={kw: "interest", "date": "ts"})
                ser["keyword"] = kw
                # ensure timestamps are JSON serializable
                ser["ts"] = ser["ts"].astype(str)
                iot_map[kw] = ser.to_dict(orient="records")
            rq = pytrends.related_queries()
            if kw in rq and rq[kw].get("rising") is not None:
                rising_df = rq[kw]["rising"].head(10)
                for _, row in rising_df.iterrows():
                    rising_all.append({
                        "keyword": kw,
                        "query": row.get("query"),
                        "value": int(row.get("value", 0) or 0),
                        "link": f"https://www.google.com/search?q={str(row.get('query','')).replace(' ', '+')}",
                    })
        except Exception as e:
            rising_all.append({"keyword": kw, "query": None, "value": 0, "error": str(e)})
    return {"source": "google_trends", "rising": rising_all, "iot": iot_map}


def reddit_hot_or_top(subreddits: List[str], mode: str = "hot", limit: int = 15) -> Dict:
    """
    Uses PRAW if keys exist; otherwise uses public JSON with a resilient UA and raw_json=1.
    Displays error reason in `note` when blocked.
    """
    try:
        import praw
    except Exception:
        # fallback to public endpoint
        praw = None

    cid = _env("REDDIT_CLIENT_ID")
    csec = _env("REDDIT_CLIENT_SECRET")
    ua = _env("REDDIT_USER_AGENT", "Wave/1.0 (https://wave; contact: owner@example.com)")

    posts = []
    note = ""
    if praw and cid and csec and ua:
        try:
            reddit = praw.Reddit(client_id=cid, client_secret=csec, user_agent=ua, check_for_async=False)
            reddit.read_only = True
            for sub in (subreddits or [])[:6]:
                try:
                    sr = reddit.subreddit(sub)
                    it = sr.top(limit=limit, time_filter="week") if mode == "top" else sr.hot(limit=limit)
                    for p in it:
                        posts.append({
                            "subreddit": sub,
                            "title": getattr(p, "title", None),
                            "score": int(getattr(p, "score", 0) or 0),
                            "url": f"https://www.reddit.com{getattr(p, 'permalink', '')}",
                            "created_utc": int(getattr(p, "created_utc", 0) or 0)
                        })
                except Exception as e:
                    note = f"PRAW error: {e}"
            posts = sorted(posts, key=lambda x: x.get("score", 0), reverse=True)
            return {"source": "reddit", "posts": posts, "note": note}
        except Exception as e:
            note = f"PRAW unavailable: {e}"

    # public endpoint fallback
    try:
        headers = {"User-Agent": ua}
        for sub in (subreddits or [])[:6]:
            url = f"https://www.reddit.com/r/{sub}/{ 'top' if mode=='top' else 'hot' }.json?raw_json=1&t=week&limit={limit}"
            r = requests.get(url, headers=headers, timeout=15)
            if r.status_code == 429:
                time.sleep(1.2)
                r = requests.get(url, headers=headers, timeout=15)
            if r.status_code == 200:
                data = r.json()
                for child in (data.get("data", {}).get("children", []) or []):
                    d = child.get("data", {})
                    posts.append({
                        "subreddit": sub,
                        "title": d.get("title"),
                        "score": int(d.get("score", 0) or 0),
                        "url": f"https://www.reddit.com{d.get('permalink','')}",
                        "created_utc": int(d.get("created_utc", 0) or 0)
                    })
            else:
                note = f"HTTP {r.status_code} from public endpoint"
        posts = sorted(posts, key=lambda x: x.get("score", 0), reverse=True)
        return {"source": "reddit", "posts": posts, "note": note}
    except Exception as e:
        return {"source": "reddit", "posts": [], "note": f"fallback error: {e}"}


def youtube_search(api_key: str, query: str, max_results: int = 10) -> Dict:
    if not api_key:
        return {"source": "youtube", "items": [], "error": "missing key"}
    try:
        url = "https://www.googleapis.com/youtube/v3/search"
        params = {"part": "snippet", "q": query, "type": "video", "order": "date",
                  "maxResults": max_results, "key": api_key}
        r = requests.get(url, params=params, timeout=20)
        r.raise_for_status()
        items = []
        for it in r.json().get("items", []):
            sn = it.get("snippet", {})
            items.append({
                "title": sn.get("title"),
                "channel": sn.get("channelTitle"),
                "publishedAt": sn.get("publishedAt"),
                "url": f"https://www.youtube.com/watch?v={it['id'].get('videoId')}"
            })
        return {"source": "youtube", "items": items}
    except Exception as e:
        return {"source": "youtube", "items": [], "error": str(e)}


def gather_trends(niche_keywords: List[str], city="", state="", subs: Optional[List[str]] = None,
                  geo="US", timeframe="now 7-d", youtube_api_key: Optional[str] = None,
                  reddit_mode: str = "hot") -> Dict:
    subs = subs or ["SmallBusiness", "Marketing"]
    kw = [k for k in niche_keywords if k][:6]
    if city and state:
        kw.append(f"{city} {state}")
    gt = google_trends_rising(kw, geo=geo, timeframe=timeframe)
    rd = reddit_hot_or_top(subs, mode=reddit_mode)
    yt = youtube_search(youtube_api_key or _env("YOUTUBE_API_KEY", ""), " | ".join(kw) or "local business", 10)
    return {
        "generated_at": dt.datetime.utcnow().isoformat(),
        "inputs": {"keywords": kw, "geo": geo, "timeframe": timeframe, "city": city, "state": state, "subs": subs},
        "google_trends": gt,
        "reddit": rd,
        "youtube": yt
    }


# ------------------ Google Places (v1 Text Search) ----------------
def search_places_optional(query: str, city: str, state: str, limit: int = 12, api_key: str = "") -> Optional[pd.DataFrame]:
    if not api_key:
        return None
    text_url = "https://places.googleapis.com/v1/places:searchText"
    loc = ", ".join([x for x in [city.strip(), state.strip()] if x])
    q = f"{query} in {loc}" if loc else query
    headers = {
        "Content-Type": "application/json",
        "X-Goog-Api-Key": api_key,
        "X-Goog-FieldMask": (
            "places.id,places.displayName,places.formattedAddress,places.location,"
            "places.rating,places.userRatingCount,places.nationalPhoneNumber,"
            "places.websiteUri,places.googleMapsUri"
        ),
    }
    body = {"textQuery": q, "maxResultCount": limit}
    r = requests.post(text_url, headers=headers, json=body, timeout=20)
    r.raise_for_status()
    places = r.json().get("places", []) or []
    rows = []
    for p in places[:limit]:
        name = (p.get("displayName") or {}).get("text") or p.get("name", "").split("/")[-1]
        locd = p.get("location") or {}
        rows.append({
            "Name": name,
            "Rating": float(p.get("rating", 0) or 0),
            "Reviews": int(p.get("userRatingCount", 0) or 0),
            "Phone": p.get("nationalPhoneNumber", ""),
            "Website": p.get("websiteUri", ""),
            "Address": p.get("formattedAddress", ""),
            "Lat": locd.get("latitude"),
            "Lng": locd.get("longitude"),
            "MapsUrl": p.get("googleMapsUri", ""),
        })
    df = pd.DataFrame(rows)
    if not df.empty:
        df["__quality"] = df["Rating"].fillna(0) * (1 + (df["Reviews"].fillna(0) / 1000))
        df = df.sort_values("__quality", ascending=False).drop(columns="__quality")
    return df


# ----------------- Streamlit app setup -------------------
st.set_page_config(page_title="Wave — AI Growth Team (Pro)", page_icon="🌊", layout="wide")
inject_css()

# Sidebar: theme + Warm-up + Pro toggles + Clear-all
st.sidebar.markdown("### Appearance")
theme = st.sidebar.radio("Theme", ["Dark", "Light"], index=0)
if theme == "Light":
    st.markdown("<style>body{background:#f6f7fb;color:#111}</style>", unsafe_allow_html=True)

st.sidebar.markdown("### Keep service awake")
wake_url = _env("RENDER_WAKE_URL", "")
if st.sidebar.button("🔥 Warm up the AI"):
    msg = warm_up_render(wake_url)
    (st.sidebar.success if msg.startswith("Warmed") else st.sidebar.warning)(msg)

st.sidebar.markdown("### Pro toggles (optional)")
use_langchain = st.sidebar.toggle("LangChain Enricher", value=False)
use_langgraph = st.sidebar.toggle("LangGraph Orchestrator", value=False)
use_crewai = st.sidebar.toggle("CrewAI Growth Crew", value=False)
autogpt_url = st.sidebar.text_input("AutoGPT Webhook URL (optional)", _env("AUTOGPT_URL", ""))

st.sidebar.markdown("---")
if st.sidebar.button("🧹 Clear all context", help="Reset picks & caches (keeps API keys)."):
    for k in [
        "trend_data_cache", "lead_data_cache", "reddit_mode",
        "niche_keywords_list", "subreddits_list", "feeder_cats_list",
        "trend_niche", "trend_subs",
        "out_persona", "industry"
    ]:
        if k in st.session_state:
            del st.session_state[k]
    st.sidebar.success("Cleared app context.")

# Session defaults
st.session_state.setdefault("trend_data_cache", None)
st.session_state.setdefault("lead_data_cache", None)
st.session_state.setdefault("reddit_mode", "hot")
st.session_state.setdefault("out_persona", "Local Professional")
st.session_state.setdefault("industry", "")

# Lists for chips
st.session_state.setdefault("niche_keywords_list", [])
st.session_state.setdefault("subreddits_list", [])
st.session_state.setdefault("feeder_cats_list", [])

st.title("🌊 Wave — AI Growth Team (Pro)")
st.caption("Trends → Leads → Outreach (+ optional LangChain, LangGraph, CrewAI).")

tab1, tab2, tab3, tab4, tab5 = st.tabs(
    ["Trend Rider", "Lead Finder", "Outreach Factory", "Weekly Report", "Pro Lab"]
)

# ===================== TREND RIDER =====================
with tab1:
    helper_banner()
    st.subheader("Trend Rider — ride what's hot")
    with st.expander("What this does", expanded=True):
        st.markdown(
            "- **Google Trends** rising queries around your niche\n"
            "- **Reddit** hot/top threads for topic hooks (real-only; AI summary if blocked)\n"
            "- **YouTube** fresh videos for zeitgeist\n"
            "- An **AI market summary** and post ideas\n"
        )

    # Industry anchor (NEW prominent field)
    c_ind, c_rank = st.columns([3,1])
    with c_ind:
        st.session_state.industry = st.text_input(
            "Industry (anchor for suggestions across the app)",
            st.session_state.get("industry", "Real Estate"),
            key="industry",
            help="This guides suggestions everywhere. You can also mix industries, comma-separated."
        )
    with c_rank:
        st.toggle("AI Rank by Opportunity", value=False, key="rank_by_opportunity", help="(future use)")

    st.selectbox("Reddit ranking", ["hot", "top"], index=0, key="reddit_mode")

    # Form
    trend_form = st.form(key="trend_form", clear_on_submit=False)
    with trend_form:
        # inputs
        c1, c2, c3 = st.columns([3,1.2,1])
        with c1:
            niche = st.text_input(
                "Niche keywords (comma-separated)",
                value=st.session_state.get("trend_niche", "real estate, mortgage, school districts"),
                key="trend_niche",
                help="Click chips below to add; you can also type here. This powers Trends/Reddit/YouTube."
            )
        with c2:
            city = st.text_input("City", "Katy", key="trend_city")
        with c3:
            state = st.text_input("State", "TX", key="trend_state")

        c4, c5 = st.columns([3,3])
        with c4:
            subs = st.text_input(
                "Reddit subs (comma-separated)",
                st.session_state.get("trend_subs", "RealEstate, Austin, personalfinance"),
                key="trend_subs",
                help="Click chips below to add; you can also type here."
            )
        with c5:
            timeframe = st.selectbox(
                "Google Trends timeframe",
                ["now 7-d", "now 1-d", "now 30-d", "today 3-m"],
                index=0,
                key="trend_timeframe",
            )

        # Suggestion buttons (AI)
        colA, colB, colC = st.columns(3)
        with colA:
            if st.button("💡 Suggest keywords", key="btn_suggest_kw"):
                prompt = (
                    f"Industry: {st.session_state.industry}\n"
                    f"Location: {city}, {state}\n"
                    f"Suggest 12 high-intent niche keywords (5-30 chars each). "
                    f"Return a JSON array of strings only."
                )
                raw = llm(prompt, system="You are a local SEO researcher.", temp=0.2)
                kws = extract_strings(raw, max_items=24)
                st.session_state["suggested_keywords"] = kws
                st.success(f"Suggested {len(kws)} keywords.")

        with colB:
            if st.button("💬 Suggest subreddits", key="btn_suggest_subs"):
                prompt = (
                    f"Industry: {st.session_state.industry}\n"
                    f"City/State: {city}, {state}\n"
                    f"Suggest 10 relevant subreddits (names only, no r/ prefix). "
                    f"Return a JSON array of strings."
                )
                raw = llm(prompt, system="You find relevant subreddits.", temp=0.3)
                subs_sug = extract_strings(raw, max_items=20)
                st.session_state["suggested_subs"] = subs_sug
                st.success(f"Suggested {len(subs_sug)} subreddits.")

        with colC:
            if st.button("🧲 Suggest feeder categories (for Lead Finder)", key="btn_suggest_feeders"):
                prompt = (
                    f"Industry: {st.session_state.industry}\n"
                    f"City/State: {city}, {state}\n"
                    f"Suggest 12 feeder business categories that see this industry's customers first "
                    f"(e.g., for real estate: apartment complexes, movers, mortgage brokers). "
                    f"Return a JSON array of short strings."
                )
                raw = llm(prompt, system="You suggest feeder businesses for partnerships.", temp=0.3)
                feeders = extract_strings(raw, max_items=24)
                st.session_state["suggested_feeders"] = feeders
                st.success(f"Suggested {len(feeders)} feeder categories.")

        # Render suggestion chips
        kw_sug = st.session_state.get("suggested_keywords", [])
        if kw_sug:
            st.success(f"Suggested {len(kw_sug)} keywords.")
            render_chips_and_capture(
                "Suggested keywords:",
                kw_sug, "kw",
                "niche_keywords_list",
                "trend_niche"
            )

        subs_sug = st.session_state.get("suggested_subs", [])
        if subs_sug:
            st.info(f"Suggested {len(subs_sug)} subreddits.")
            render_chips_and_capture(
                "Suggested subreddits:",
                subs_sug, "sub",
                "subreddits_list",
                "trend_subs"
            )

        feed_sug = st.session_state.get("suggested_feeders", [])
        if feed_sug:
            st.warning(f"Suggested {len(feed_sug)} feeder categories.")
            render_chips_and_capture(
                "Suggested feeder categories:",
                feed_sug, "feed",
                "feeder_cats_list",
                "lead_cat"  # will also prefill on Lead Finder
            )

        submitted = st.form_submit_button("Fetch trends", key="trend_submit")

    # Compute on submit
    if submitted:
        # sync typed inputs into lists (supports manual typing)
        st.session_state["niche_keywords_list"] = [s.strip() for s in st.session_state.trend_niche.split(",") if s.strip()]
        st.session_state["subreddits_list"] = [s.strip().lstrip("r/") for s in st.session_state.trend_subs.split(",") if s.strip()]

        keywords = st.session_state["niche_keywords_list"]
        sub_list = st.session_state["subreddits_list"]
        data = gather_trends(
            niche_keywords=keywords, city=city, state=state, subs=sub_list,
            timeframe=timeframe, youtube_api_key=_env("YOUTUBE_API_KEY", ""),
            reddit_mode=st.session_state.reddit_mode
        )
        st.session_state.trend_data_cache = data

    data = st.session_state.trend_data_cache
    if not data:
        st.info("Enter your niche/city and click **Fetch trends**.")
    else:
        rising = pd.DataFrame(data.get("google_trends", {}).get("rising", []))
        st.markdown("### Google Trends — Rising Queries")
        if not rising.empty:
            st.dataframe(rising[["keyword", "query", "value", "link"]], use_container_width=True)
        else:
            st.info("No rising queries or pytrends missing.")

        st.markdown("### Reddit — Hot Posts")
        rd = data.get("reddit", {})
        note = rd.get("note", "")
        if note:
            st.caption(f"Reddit note: {note}")
        posts = pd.DataFrame(rd.get("posts", []))
        if not posts.empty:
            st.dataframe(posts[["subreddit", "title", "score", "url"]].head(20), use_container_width=True)
        else:
            st.warning("Reddit posts unavailable (API blocked or no data).")
            # AI fallback summary (not posts)
            sample_kw = ", ".join(st.session_state.get("niche_keywords_list", [])[:5])
            ai_box = llm(
                system="You summarize discussion themes (not fabricating posts).",
                prompt=f"Summarize likely discussion themes around: {sample_kw} in {city}, {state}. "
                       f"Return 5 bullets that could inspire content ideas."
            ) or "AI unavailable."
            st.info("**Reddit AI fallback (insights, not real posts):**\n\n" + ai_box)

        st.markdown("### YouTube — Fresh Videos (optional)")
        yt = data.get("youtube", {})
        vids = pd.DataFrame(yt.get("items", []))
        if not vids.empty:
            st.dataframe(vids[["title", "channel", "publishedAt", "url"]], use_container_width=True)
        else:
            if yt.get("error"):
                st.caption(f"YouTube: {yt['error']}")
            else:
                st.caption("Add YOUTUBE_API_KEY to show videos.")

        st.markdown("### AI Market Summary")
        sample = {
            "trending_queries": rising.head(8).to_dict(orient="records") if not rising.empty else [],
            "reddit_top": posts.head(8).to_dict(orient="records") if not posts.empty else [],
        }
        summary = llm(
            system="You are a concise SMB strategist.",
            prompt=(f"Summarize ~5 bullets of what's trending for {city}, {state} in industry {st.session_state.industry}. "
                    f"Then propose 3 ride-the-wave post ideas. Data:\n{json.dumps(sample)}")
        ) or "Add OPENAI_API_KEY to enable AI-written summaries."
        st.info(summary)


# ===================== LEAD FINDER =====================
with tab2:
    helper_banner()
    st.subheader("Lead Finder — Nearby partners & feeder businesses")
    with st.expander("How this helps", expanded=True):
        st.markdown(
            "- We surface **feeder businesses** that meet your future customers first.\n"
            "- You get **phone, website, directions** to start a partnership.\n"
            "- We compute an **Actionability Score** and explain **why** each lead is hot.\n"
            "- Typical targets for real estate: *apartment complex, movers, mortgage broker, home builder*.\n"
        )

    # Prefill from chips if present
    default_cat = st.session_state.get("lead_cat", "apartment complex")
    lead_form = st.form(key="lead_form", clear_on_submit=False)
    with lead_form:
        cat = st.text_input("Place type / query", default_cat, key="lead_cat", help="Try any feeder category. You can also click suggestions in Trend Rider.")
        city2 = st.text_input("City", st.session_state.get("trend_city", "Katy"), key="lead_city")
        state2 = st.text_input("State", st.session_state.get("trend_state", "TX"), key="lead_state")
        limit = st.slider("How many?", 5, 30, 12, key="lead_limit")
        go = lead_form.form_submit_button("Search", key="lead_submit")

    def actionability_score(row, query: str):
        score, reasons = 0, []
        r = float(row.get("Rating") or 0)
        n = int(row.get("Reviews") or 0)
        if r >= 4.4:
            score += 35; reasons.append("High rating (≥4.4)")
        elif r >= 4.0:
            score += 20; reasons.append("Solid rating (≥4.0)")
        if n >= 200:
            score += 35; reasons.append("Strong review volume (≥200)")
        elif n >= 50:
            score += 20; reasons.append("Decent review volume (≥50)")
        if "apartment" in query.lower():
            score += 20; reasons.append("Feeder: high tenant churn → frequent moves")
        if row.get("Website"):
            score += 10; reasons.append("Website present")
        if row.get("Phone"):
            score += 10; reasons.append("Phone present")
        return min(score, 100), reasons

    if go:
        base_df = search_places_optional(cat, city2, state2, limit=limit, api_key=_env("GOOGLE_PLACES_API_KEY", ""))
        st.session_state.lead_data_cache = base_df if base_df is not None else False

    df = st.session_state.lead_data_cache
    if df is None:
        st.info("Enter a query (e.g., **apartment complex**, **moving company**, **mortgage broker**, **home builder**).")
    elif df is False or (isinstance(df, pd.DataFrame) and df.empty):
        st.warning("No results (check GOOGLE_PLACES_API_KEY or try a nearby city/another query).")
    else:
        scored = df.copy()
        scored["Score"] = 0
        scored["Why"] = ""
        for i, row in scored.iterrows():
            s, rs = actionability_score(row, cat)
            scored.at[i, "Score"] = int(s)
            scored.at[i, "Why"] = " · ".join(rs)

        c1, c2, c3 = st.columns(3)
        with c1: kpi("Shown", str(len(scored)))
        with c2: kpi("Avg rating", f"{scored['Rating'].mean():.2f}" if len(scored) else "–")
        with c3: kpi("Median reviews", f"{int(scored['Reviews'].median())}" if len(scored) else "–")

        # table
        dff = scored.reset_index(drop=True).copy()
        st.markdown("#### Ranked leads")
        show_cols = ["Name", "Score", "Why", "Rating", "Reviews", "Phone", "Website", "Address"]
        st.dataframe(dff[show_cols], use_container_width=True)

        # Select + details panel
        if "Name" in dff.columns and len(dff) > 0:
            name_choice = st.selectbox("Select a business", list(dff["Name"].astype(str).unique()))
        else:
            st.warning("No 'Name' column in the results.")
            name_choice = None

        if name_choice:
            row = dff[dff["Name"].astype(str) == str(name_choice)].iloc[0].to_dict()
            st.markdown(f"##### 📍 {row.get('Name','')}")
            st.markdown(f"{row.get('Address','')}")
            st.markdown(
                f"- ⭐ **{row.get('Rating','—')}** ({row.get('Reviews','—')} reviews)\n"
                f"- 📞 {row.get('Phone','—')}\n"
                f"- 🌐 {row.get('Website','—')}\n"
                f"- 🧭 [Open in Google Maps]({row.get('MapsUrl','')})"
            )

        # map
        if MAPS_OK:
            st.markdown("#### Map")
            dfc = dff.dropna(subset=["Lat", "Lng"]).copy()
            if not dfc.empty:
                m = folium.Map(location=[dfc["Lat"].mean(), dfc["Lng"].mean()], zoom_start=12)
                for _, r in dfc.iterrows():
                    folium.Marker(
                        [r["Lat"], r["Lng"]],
                        tooltip=r["Name"],
                        popup=f"<b>{r['Name']}</b><br>{r['Address']}<br>{r['Rating']} ⭐ / {r['Reviews']} reviews<br>"
                              f"<a href='{r['MapsUrl']}' target='_blank'>Open in Google Maps</a>",
                    ).add_to(m)
                st_folium(m, height=520)
            else:
                st.info("No coordinates to plot.")
        else:
            st.caption("Install folium + streamlit-folium to see the map.")

        st.download_button("⬇️ Export CSV", dff[show_cols].to_csv(index=False).encode("utf-8"),
                           "leads.csv", "text/csv")

        # One-click outreach
        st.markdown("#### One-click outreach")
        if name_choice:
            lead2 = dff[dff["Name"].astype(str) == str(name_choice)].iloc[0].to_dict()
            colA, colB = st.columns(2)
            with colA:
                if st.button("Draft email", use_container_width=True, key="btn_email"):
                    body = (llm(
                        system="You write friendly B2B outreach for local partnerships.",
                        prompt=(f"Draft a short email from a {st.session_state.get('out_persona','Local Professional')} "
                                f"to {lead2['Name']} ({lead2.get('Website','')}). "
                                f"Goal: propose a referral partnership. Include a clear CTA. "
                                f"Keep it 120–150 words.")
                    ) or
                    f"Hi {lead2['Name']} team,\n\nI’d love to explore a simple referral partnership. "
                    "We serve the same customers and can help each other win more business. "
                    "Could we schedule a 10-minute chat this week?\n\nBest,\n<Your Name>")
                    st.code(body)
            with colB:
                if st.button("Draft SMS", use_container_width=True, key="btn_sms"):
                    sms = (llm(
                        system="You write concise SMS for local B2B outreach.",
                        prompt=(f"Write a friendly 240-character text to {lead2['Name']} proposing a quick chat about referrals. "
                                f"One clear CTA.")
                    ) or
                    f"Hi {lead2['Name']}! Quick idea: partner on referrals? 10-min chat this week?")
                    st.code(sms)

        # AutoGPT webhook (optional)
        if autogpt_url:
            if st.button("Arm AutoGPT: watch for new high-score leads", key="btn_autogpt"):
                payload = {"action": "lead_watch", "query": cat, "city": city2, "state": state2, "threshold": 85}
                try:
                    r = requests.post(autogpt_url, json=payload, timeout=15)
                    st.success(f"Webhook sent: {r.status_code}")
                except Exception as e:
                    st.warning(f"Webhook failed: {e}")


# ===================== OUTREACH FACTORY =====================
with tab3:
    helper_banner()
    st.subheader("Outreach Factory — ready-to-send sequences")
    st.caption("AI-polished emails/SMS with fixed send dates. Export as TXT/DOCX.")

    c1, c2 = st.columns(2)
    with c1:
        persona = st.text_input("Your business type", st.session_state.get("out_persona", "Real Estate Agent"),
                                key="out_persona")
        target = st.text_input("Target audience", "Apartment complex managers in Katy, TX", key="out_target")
    with c2:
        tone = st.selectbox("Tone", ["Friendly", "Professional", "Hype"], index=0, key="out_tone")
        touches = st.slider("Number of touches", 3, 6, 3, key="out_touches")

    if st.button("Generate Sequence", key="out_generate"):
        base = [
            {"send_dt": str(dt.date.today()), "channel": "email", "subject": "Quick hello 👋",
             "body": f"Hi there — I’m a {persona} in {target}. Could we collaborate?"},
            {"send_dt": str(dt.date.today() + dt.timedelta(days=2)), "channel": "sms", "subject": "",
             "body": "Hey! Just checking in — open to a quick chat this week?"},
            {"send_dt": str(dt.date.today() + dt.timedelta(days=7)), "channel": "email", "subject": "Ready when you are",
             "body": "Happy to help with referrals and co-marketing. What works?"}
        ][:touches]

        prompt = (
            f"Polish this outreach for a {persona} to contact {target}. "
            f"Keep SAME dates and channels. Tone: {tone}. Return PLAIN TEXT (not JSON). "
            f"Here are the steps as JSON for reference:\n{json.dumps(base)}"
        )
        polished = llm(prompt, system="You write high-converting SMB outreach.") or \
                   "\n".join([f"{s['send_dt']} • {s['channel']}: {s['subject']} {s['body']}".strip() for s in base])

        st.markdown("### AI-Polished Copy")
        st.markdown(polished)

        st.download_button("⬇️ Download TXT", polished.encode("utf-8"), "outreach.txt", "text/plain")
        if DOCX_OK:
            st.download_button("⬇️ Download DOCX",
                build_docx_bytes("Outreach Plan", polished),
                "outreach.docx",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        else:
            st.caption("Install `python-docx` for DOCX export.")


# ===================== WEEKLY REPORT =====================
with tab4:
    helper_banner()
    st.subheader("Weekly Output (PDF/DOCX-ready text)")
    st.caption("Combines Trend hooks + Leads + Outreach into a single narrative.")

    biz = st.text_input("Business type", st.session_state.get("out_persona", "Real Estate Agent"), key="report_biz")
    rcity = st.text_input("City for report", st.session_state.get("trend_city", "Katy"), key="report_city")
    date = st.date_input("Week of", dt.date.today(), key="report_date")

    if st.button("Build Weekly Report", key="report_build"):
        report = textwrap.dedent(f"""
        # Wave Weekly Report — {biz}, {rcity}
        **Week of:** {date}

        ## 1) Lead Finder — where your next clients are
        - Pull feeder businesses (apartments, movers, mortgage brokers).
        - Sort by **Actionability Score**; call high-scoring leads first.
        - Use one-click outreach for fast wins.

        ## 2) Trend Rider — ride what's hot
        - Look for rising searches (Google Trends) & hot Reddit threads.
        - Publish 3 posts that hit the trending questions this week.

        ## 3) Outreach Factory — send this today
        - 3-touch sequence ready (TXT/DOCX).
        - Keep it friendly, concrete, and localized.

        _Wave — AI Growth Team_
        """).strip()

        st.text(report)
        st.download_button("⬇️ Download TXT", report.encode("utf-8"), "weekly_report.txt", "text/plain")
        if DOCX_OK:
            st.download_button("⬇️ Download DOCX",
                build_docx_bytes("Weekly Report", report),
                "weekly_report.docx",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document")


# ===================== PRO LAB (stubs) =====================
with tab5:
    helper_banner()
    st.subheader("Pro Lab — LangChain · LangGraph · CrewAI (optional)")
    st.caption("These automate deeper research/polish. The app runs fine without them.")

    # Diagnostics
    with st.expander("Pro diagnostics", expanded=False):
        diag = {
            "LangChain core (LC_OK)": LC_OK,
            "LangChain OpenAI (LCO_OK)": LCO_OK,
            "LangGraph (LG_OK)": LG_OK,
            "CrewAI (CREW_OK)": CREW_OK,
            "OPENAI_API_KEY present": bool(_env("OPENAI_API_KEY", "")),
            "YOUTUBE_API_KEY present": bool(_env("YOUTUBE_API_KEY", "")),
            "GOOGLE_PLACES_API_KEY present": bool(_env("GOOGLE_PLACES_API_KEY", "")),
            "Last Reddit source": st.session_state.get("reddit_mode",""),
        }
        st.json(diag)

    # LangChain Enricher
    st.markdown("### LangChain Enricher")
    if not use_langchain:
        st.info("Toggle **LangChain Enricher** in the sidebar to enable.")
    elif not (LC_OK and LCO_OK and OPENAI_API_KEY):
        st.warning("LangChain libraries or OPENAI_API_KEY missing.")
    else:
        leads_df = st.session_state.lead_data_cache
        if isinstance(leads_df, pd.DataFrame) and not leads_df.empty:
            top_n = st.slider("How many leads to enrich?", 3, 15, 5, key="lc_topn")
            model = ChatOpenAI(model="gpt-4o-mini", temperature=0.2, api_key=OPENAI_API_KEY)
            prompt = ChatPromptTemplate.from_template(
                "Rate partner fit for referrals.\nLead: {lead}\nReturn JSON: fit(1-100), why, next_action."
            )
            enriched_rows = []
            for _, row in leads_df.head(top_n).iterrows():
                lead_json = row.to_dict()
                try:
                    msg = model.invoke(prompt.format_messages(lead=json.dumps(lead_json)))
                    txt = (msg.content or "").strip()
                    data = {}
                    try:
                        data = json.loads(txt)
                    except Exception:
                        start = txt.find("{"); end = txt.rfind("}")
                        if start != -1 and end != -1:
                            data = json.loads(txt[start:end+1])
                    enriched_rows.append({
                        "Name": row["Name"],
                        "LC_Fit": data.get("fit", 50),
                        "LC_Why": data.get("why", ""),
                        "LC_Next": data.get("next_action", ""),
                    })
                except Exception as e:
                    enriched_rows.append({"Name": row["Name"], "LC_Fit": 50, "LC_Why": f"error: {e}", "LC_Next": ""})
            st.dataframe(pd.DataFrame(enriched_rows), use_container_width=True)
        else:
            st.info("Run **Lead Finder** first so we have leads to enrich.")

    st.divider()

    # LangGraph Orchestrator
    st.markdown("### LangGraph Orchestrator")
    if not use_langgraph:
        st.info("Toggle **LangGraph Orchestrator** in the sidebar to enable.")
    elif not (LG_OK and LCO_OK and OPENAI_API_KEY):
        st.warning("LangGraph libraries or OPENAI_API_KEY missing.")
    else:
        trend_data = st.session_state.trend_data_cache
        if not trend_data:
            st.info("Fetch trends in **Trend Rider** first.")
        else:
            from typing import TypedDict
            class S(TypedDict):
                intent: str
                brief: str
            model = ChatOpenAI(model="gpt-4o-mini", temperature=0.2, api_key=OPENAI_API_KEY)
            def decide(state: S) -> S:
                rising = trend_data.get("google_trends", {}).get("rising", [])
                hot = len([r for r in rising if r.get("value", 0) >= 100]) >= 3
                state["intent"] = "content" if hot else "research"
                return state
            def do_research(state: S) -> S:
                msg = model.invoke(f"Summarize 5 bullet insights from:\n{json.dumps(trend_data)[:6000]}")
                state["brief"] = (msg.content or "").strip()
                return state
            def do_content(state: S) -> S:
                msg = model.invoke("Draft 3 short post ideas with hooks + CTAs for a local business based on these trends:\n"
                                   + json.dumps(trend_data)[:6000])
                state["brief"] = (msg.content or "").strip()
                return state
            g = StateGraph(S)
            g.add_node("decide", decide)
            g.add_node("research", do_research)
            g.add_node("content", do_content)
            g.set_entry_point("decide")
            def router(state: S):
                return state["intent"]
            g.add_conditional_edges("decide", router, {"research": "research", "content": "content"})
            g.add_edge("research", END)
            g.add_edge("content", END)
            app = g.compile()
            out = app.invoke({"intent": "", "brief": ""})
            st.info(f"Intent: **{out['intent']}**")
            st.markdown(out["brief"])

    st.divider()

    # CrewAI Growth Crew
    st.markdown("### CrewAI Growth Crew")
    if not use_crewai:
        st.info("Toggle **CrewAI Growth Crew** in the sidebar to enable.")
    elif not (CREW_OK and OPENAI_API_KEY):
        st.warning("CrewAI not installed or OPENAI_API_KEY missing.")
    else:
        biz = st.text_input("Business", st.session_state.get("out_persona", "Real Estate Agent"), key="crew_biz")
        niche = st.text_input("Niche focus", "Relocation in Katy, TX", key="crew_niche")
        if st.button("Run Growth Crew", key="crew_run"):
            researcher = Agent(
                role="Market Researcher",
                goal="Find insights & angles that will convert for a local SMB.",
                backstory="You love concise facts and actionable takeaways.",
                allow_delegation=False,
                verbose=False,
                llm=ChatOpenAI(model="gpt-4o-mini", temperature=0.1, api_key=OPENAI_API_KEY),
            )
            writer = Agent(
                role="Copywriter",
                goal="Write a crisp one-pager a business owner can use today.",
                backstory="You are crisp, persuasive, and practical.",
                allow_delegation=False,
                verbose=False,
                llm=ChatOpenAI(model="gpt-4o-mini", temperature=0.3, api_key=OPENAI_API_KEY),
            )
            t1 = Task(description=f"Summarize 5 bullet insights for {biz} ({niche}). Provide sources if relevant.", agent=researcher)
            t2 = Task(description="Write a 1-page brief with headline, 3 bullets, 1 CTA. Make it print-friendly.", agent=writer)
            crew = Crew(agents=[researcher, writer], tasks=[t1, t2])
            try:
                result = crew.kickoff()
            except Exception as e:
                result = f"(CrewAI runtime error: {e})"
            st.markdown("#### One-pager")
            st.markdown(str(result))

# If OPENAI_API_KEY is missing, show a gentle hint (does not stop the app)
if client is None and not _env("OPENAI_API_KEY", ""):
    st.caption("Set OPENAI_API_KEY in your environment or Streamlit secrets to enable AI features.")
